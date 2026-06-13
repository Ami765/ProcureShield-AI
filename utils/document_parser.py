"""
utils/document_parser.py — Document Parser Utility
Extracts plain text from uploaded PDF or DOCX files.
"""

import io
from typing import Optional


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract all text from a PDF file given its raw bytes."""
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n\n".join(pages).strip()
        if not text:
            return "[PDF contained no extractable text — it may be a scanned image.]"
        return text
    except Exception as e:
        return f"[Error reading PDF: {e}]"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract all text from a DOCX file given its raw bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n\n".join(paragraphs).strip() or "[DOCX contained no extractable text.]"
    except Exception as e:
        return f"[Error reading DOCX: {e}]"


def parse_uploaded_file(file_name: str, file_bytes: bytes) -> str:
    """
    Dispatch to the correct parser based on file extension.

    Parameters
    ----------
    file_name  : Original filename (used to detect extension).
    file_bytes : Raw file content as bytes.

    Returns
    -------
    Extracted plain text string.
    """
    ext = file_name.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    else:
        # Attempt plain-text decode as fallback
        try:
            return file_bytes.decode("utf-8", errors="replace")
        except Exception:
            return "[Unsupported file format.]"
